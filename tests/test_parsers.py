"""解析器测试（FR-201 / FR-202 / FR-203）。"""

from __future__ import annotations

import math

import pytest
from docx import Document as DocxDocument
from reportlab.pdfgen import canvas as pdf_canvas

from docgraph.parsers import DocxParser, PdfParser, get_parser
from docgraph.parsers.base import ScannedPdfError, build_chunks, estimate_tokens


def _make_pdf(path: str, texts: list[str]) -> None:
    """用 reportlab 生成含指定文本的 PDF（一页一段文本）。"""
    c = pdf_canvas.Canvas(str(path))
    for text in texts:
        c.drawString(72, 720, text)
        c.showPage()
    c.save()


def _make_blank_pdf(path: str) -> None:
    """生成无文本层的 PDF（模拟扫描件，FR-203）。"""
    c = pdf_canvas.Canvas(str(path))
    c.setFillColorRGB(1, 1, 1)
    c.rect(0, 0, 600, 800, stroke=0, fill=1)
    c.showPage()
    c.save()


# ---- FR-202：分块 ----

def test_estimate_tokens():
    assert estimate_tokens("abc") == 2  # 3 字符 / 2 向上取整
    assert estimate_tokens("") == 1


def test_build_chunks_splits_and_overlaps():
    sections = [(1, f"第{i}段：" + "内容" * 50) for i in range(1, 6)]  # 每段约 100 字符 ≈ 50 token
    chunks = build_chunks("doc1", sections, max_tokens=120, overlap_tokens=20)
    assert len(chunks) >= 2
    # 块序号连续、页码为块内首段页码
    for i, c in enumerate(chunks):
        assert c.seq == i
        assert c.document_id == "doc1"
        assert c.id == f"doc1:c{i}"
        assert c.token_count > 0
    # 相邻块有重叠
    assert chunks[1].text.startswith(chunks[0].text[-40:])


def test_build_chunks_ignores_empty():
    chunks = build_chunks("doc1", [(1, "  "), (2, "有内容")])
    assert len(chunks) == 1
    assert "有内容" in chunks[0].text


# ---- FR-201 / FR-203：PDF ----

def test_pdf_parser_extracts_text(tmp_path):
    pdf = tmp_path / "sample.pdf"
    # 每页文本需超过扫描件阈值（50 字符/页），避免误判（FR-203）
    _make_pdf(
        pdf,
        [
            "Graph Neural Network is a powerful model for learning on graph-structured data.",
            "Attention mechanism improves the performance of graph neural networks significantly.",
        ],
    )
    chunks = PdfParser().parse(str(pdf), "d1")
    assert chunks
    assert "Graph Neural Network" in chunks[0].text
    assert chunks[0].page == 1  # 页码保留（FR-201）


def test_pdf_parser_detects_scanned(tmp_path):
    pdf = tmp_path / "scanned.pdf"
    _make_blank_pdf(pdf)
    with pytest.raises(ScannedPdfError):
        PdfParser().parse(str(pdf), "d1")


# ---- FR-201：DOCX ----

def test_docx_parser_extracts_paragraphs_and_tables(tmp_path):
    docx_path = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_paragraph("第一篇文档内容：知识图谱构建方法。")
    doc.add_paragraph("第二段：实体与关系抽取。")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "方法"
    table.rows[0].cells[1].text = "说明"
    doc.save(str(docx_path))

    chunks = DocxParser().parse(str(docx_path), "d2")
    text = "\n".join(c.text for c in chunks)
    assert "知识图谱构建方法" in text
    assert "方法 | 说明" in text
    assert all(c.page == 0 for c in chunks)


def test_parser_registry():
    assert get_parser("pdf").format == "pdf"
    assert get_parser("docx").format == "docx"
    with pytest.raises(ValueError):
        get_parser("xyz")
