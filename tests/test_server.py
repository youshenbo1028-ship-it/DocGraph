"""服务层集成测试：导入 -> 解析 -> 抽取（假抽取器）-> 图谱查询。"""

from __future__ import annotations

import io

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient

import docgraph.server.app as app_mod
from docgraph.core.models import Entity, Relation
from docgraph.extractors.base import ExtractionResult

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class FakeExtractor:
    """固定返回 GNN -> Transformer 关系的假抽取器（不联网）。"""

    def extract(self, chunk_text, known_entities=None, chunk_id=""):
        return ExtractionResult(
            entities=[
                Entity(id="", canonical_name="GNN", type="概念/方法/理论", confidence=0.9),
                Entity(id="", canonical_name="Transformer", type="概念/方法/理论", confidence=0.8),
            ],
            relations=[
                Relation(
                    id="",
                    source_entity_id="GNN",
                    target_entity_id="Transformer",
                    type="提出",
                    evidence=["原文摘录"],
                    confidence=0.9,
                )
            ],
            chunk_id=chunk_id,
        )


@pytest.fixture
def client(tmp_path, monkeypatch):
    monkeypatch.setattr(app_mod, "DATA_DIR", tmp_path / "data")
    monkeypatch.setenv("DOCGRAPH_SETTINGS_PATH", str(tmp_path / "settings.json"))
    app_mod.registry.clear()
    monkeypatch.setattr(
        app_mod,
        "make_extractor_factory",
        lambda store, api: (lambda group_id=None: FakeExtractor()),
    )
    return TestClient(app_mod.app)


def _make_docx_bytes() -> bytes:
    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_paragraph("Graph Neural Network 是一种用于图结构数据的深度学习模型。")
    doc.add_paragraph("Transformer 架构使用注意力机制。GNN 借鉴了 Transformer 的思想。")
    doc.save(buf)
    return buf.getvalue()


def test_full_pipeline(client):
    # 1) 创建项目
    r = client.post("/api/projects", json={"name": "集成测试"})
    assert r.status_code == 200
    pid = r.json()["project"]["id"]
    gid = r.json()["groups"][0]["id"]

    # 2) 导入 DOCX
    r = client.post(
        f"/api/projects/{pid}/documents",
        files={"file": ("paper.docx", _make_docx_bytes(), DOCX_MIME)},
        data={"group_id": gid},
    )
    assert r.status_code == 200, r.text
    doc_id = r.json()["id"]

    # 3) 解析
    r = client.post(f"/api/projects/{pid}/documents/{doc_id}/parse")
    assert r.status_code == 200
    assert r.json()["status"] == "parsed"

    # 4) 抽取（假抽取器）
    r = client.post(
        f"/api/projects/{pid}/extract",
        json={"group_id": gid, "api": {"base_url": "http://x", "api_key": "k", "model": "m"}},
    )
    assert r.status_code == 200, r.text
    summary = r.json()
    assert summary["documents"] == 1
    assert summary["entities"] >= 2
    assert summary["relations"] >= 1

    # 5) 图谱查询
    r = client.get(f"/api/projects/{pid}/graph", params={"group_id": gid})
    assert r.status_code == 200
    graph = r.json()
    assert len(graph["nodes"]) >= 2
    assert len(graph["edges"]) >= 1
    labels = {n["data"]["label"] for n in graph["nodes"]}
    assert {"GNN", "Transformer"} <= labels


def test_duplicate_import_conflict(client):
    pid = client.post("/api/projects", json={"name": "t"}).json()["project"]["id"]
    gid = client.post(f"/api/projects/{pid}/groups", json={"name": "g"}).json()["id"]
    r1 = client.post(
        f"/api/projects/{pid}/documents",
        files={"file": ("a.docx", _make_docx_bytes(), DOCX_MIME)},
        data={"group_id": gid},
    )
    assert r1.status_code == 200
    r2 = client.post(
        f"/api/projects/{pid}/documents",
        files={"file": ("a.docx", _make_docx_bytes(), DOCX_MIME)},
        data={"group_id": gid},
    )
    assert r2.status_code == 409  # FR-310：同名文档拒绝


def test_unsupported_format(client):
    pid = client.post("/api/projects", json={"name": "t"}).json()["project"]["id"]
    r = client.post(f"/api/projects/{pid}/documents", files={"file": ("a.txt", b"hello", "text/plain")})
    assert r.status_code == 400  # FR-102：MVP 仅 PDF/DOCX


def test_settings_endpoints(client):
    """FR-801/802：保存并读取 API 配置（不返回 Key 明文）。"""
    r = client.post(
        "/api/settings",
        json={"base_url": "https://api.example.com/v1", "api_key": "sk-secret-1", "model": "m1"},
    )
    assert r.status_code == 200
    body = r.json()
    assert body["base_url"] == "https://api.example.com/v1"
    assert body["has_key"] is True
    assert "sk-secret-1" not in str(body)  # 不返回 Key 明文

    r2 = client.get("/api/settings")
    assert r2.json()["base_url"] == "https://api.example.com/v1"


def _make_project_with_graph(client) -> tuple[str, str]:
    pid = client.post("/api/projects", json={"name": "导出测试"}).json()["project"]["id"]
    gid = client.post(f"/api/projects/{pid}/groups", json={"name": "g"}).json()["id"]
    r = client.post(
        f"/api/projects/{pid}/documents",
        files={"file": ("p.docx", _make_docx_bytes(), DOCX_MIME)},
        data={"group_id": gid},
    )
    doc_id = r.json()["id"]
    client.post(f"/api/projects/{pid}/documents/{doc_id}/parse")
    client.post(
        f"/api/projects/{pid}/extract",
        json={"group_id": gid, "api": {"base_url": "http://x", "api_key": "k", "model": "m"}},
    )
    return pid, gid


def test_export_nodes_and_edges_csv(client):
    """FR-602：CSV 导出（可导入 Gephi）。"""
    pid, gid = _make_project_with_graph(client)
    r = client.get(f"/api/projects/{pid}/export/nodes.csv", params={"group_id": gid})
    assert r.status_code == 200
    assert "text/csv" in r.headers["content-type"]
    assert "GNN" in r.text
    assert r.text.splitlines()[0] == "id,label,type,confidence"

    r2 = client.get(f"/api/projects/{pid}/export/edges.csv", params={"group_id": gid})
    assert r2.status_code == 200
    assert "source,target,type,confidence,evidence" in r2.text.splitlines()[0]


def test_export_graph_json(client):
    """FR-602：Graph JSON 自描述导出。"""
    pid, gid = _make_project_with_graph(client)
    r = client.get(f"/api/projects/{pid}/export/graph.json", params={"group_id": gid})
    assert r.status_code == 200
    body = r.json()
    assert body["schema"] == "docgraph-graph/v1"
    assert body["stats"]["nodes"] >= 2
    assert body["stats"]["edges"] >= 1

def test_active_project_with_no_projects(client):
    """无项目时 active 返回空（前端据此新建）。"""
    r = client.get("/api/projects/active")
    assert r.status_code == 200
    assert r.json()["project"] is None


def test_activate_and_active_roundtrip(client):
    """激活项目 -> last_project_id 持久化 -> active 恢复。"""
    pid = client.post("/api/projects", json={"name": "t"}).json()["project"]["id"]
    r = client.post(f"/api/projects/{pid}/activate")
    assert r.status_code == 200
    assert r.json()["project"]["id"] == pid
    r2 = client.get("/api/projects/active")
    assert r2.json()["project"]["id"] == pid
    from docgraph.core import settings
    assert settings.get_last_project_id() == pid


def test_list_projects_includes_doc_count(client):
    pid = client.post("/api/projects", json={"name": "t"}).json()["project"]["id"]
    gid = client.post(f"/api/projects/{pid}/groups", json={"name": "g"}).json()["id"]
    client.post(
        f"/api/projects/{pid}/documents",
        files={"file": ("p.docx", _make_docx_bytes(), DOCX_MIME)},
        data={"group_id": gid},
    )
    projs = client.get("/api/projects").json()
    item = next(p for p in projs if p["id"] == pid)
    assert item["doc_count"] >= 1


def test_extract_uses_stored_key_when_request_empty(client):
    """请求里 api_key 为空 -> 后端读取已存 Key，不再返回 400（问题2 修复）。"""
    from docgraph.core import settings as settings_mod
    client.app  # noqa
    pid = client.post("/api/projects", json={"name": "t"}).json()["project"]["id"]
    gid = client.post(f"/api/projects/{pid}/groups", json={"name": "g"}).json()["id"]
    # 保存 key（强制文件回退，避免污染系统凭据库）
    settings_mod.keyring = None
    client.post("/api/settings", json={"base_url": "http://x", "api_key": "sk-test", "model": "m"})
    # api_key 为空 -> 后端应从 settings 读 key
    r = client.post(
        f"/api/projects/{pid}/extract",
        json={"group_id": gid, "api": {"base_url": "http://x", "api_key": "", "model": "m"}},
    )
    assert r.status_code == 200  # 而非 400" +

def test_extract_processes_pending_document(client):
    """待处理（pending）文档也应参与抽取（避免刚导入即被跳过）。"""
    pid = client.post("/api/projects", json={"name": "t"}).json()["project"]["id"]
    gid = client.post(f"/api/projects/{pid}/groups", json={"name": "g"}).json()["id"]
    r = client.post(
        f"/api/projects/{pid}/documents",
        files={"file": ("p.docx", _make_docx_bytes(), DOCX_MIME)},
        data={"group_id": gid},
    )
    assert r.status_code == 200
    assert r.json()["status"] == "pending"
    # 直接抽取（不先 parse）—— 流程应自动解析并抽取出实体
    resp = client.post(
        f"/api/projects/{pid}/extract",
        json={"group_id": gid, "api": {"base_url": "http://x", "api_key": "k", "model": "m"}},
    )
    assert resp.status_code == 200
    # 假抽取器会产出 GNN/Transformer => 至少 1 文档被处理
    assert resp.json()["documents"] == 1

def test_entity_and_relation_detail_evidence(client):
    """FR-307：实体/关系详情返回来源文档与原文依据。"""
    pid, gid = _make_project_with_graph(client)
    graph = client.get(f"/api/projects/{pid}/graph", params={"group_id": gid}).json()
    node = graph["nodes"][0]
    r = client.get(f"/api/projects/{pid}/entities/{node['data']['id']}")
    assert r.status_code == 200, r.text
    d = r.json()
    assert d["canonical_name"]
    assert d["source_docs"]  # 来源文档
    # 实体证据（原文摘录，从分块检索）
    assert "evidence" in d

    edge = graph["edges"][0]
    r2 = client.get(f"/api/projects/{pid}/relations/{edge['data']['id']}")
    assert r2.status_code == 200, r2.text
    rd = r2.json()
    assert rd["source"] and rd["target"] and rd["type"]
    assert rd["evidence"]  # 关系证据（原文摘录）

def test_extract_requires_api_config(client):
    pid = client.post("/api/projects", json={"name": "t"}).json()["project"]["id"]
    r = client.post(
        f"/api/projects/{pid}/extract",
        json={"group_id": None, "api": {"base_url": "", "api_key": "", "model": ""}},
    )
    assert r.status_code == 400  # FR-801：缺少 API 配置时给出引导
