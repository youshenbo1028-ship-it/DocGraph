"""Word（DOCX）解析器（FR-102 / FR-201）。"""

from __future__ import annotations

import docx

from ..core.models import Chunk
from .base import DocumentParser, build_chunks


class DocxParser(DocumentParser):
    format = "docx"

    def parse(self, path: str, document_id: str = "") -> list[Chunk]:
        """提取段落与表格文本（DOCX 无页码，统一 page=0）并按章节分块。"""
        d = docx.Document(path)
        sections: list[tuple[int, str]] = []

        for para in d.paragraphs:
            text = para.text.strip()
            if text:
                sections.append((0, text))

        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    sections.append((0, " | ".join(cells)))

        return build_chunks(document_id, sections)
