"""生成样例文档（examples/）：用于本地体验与测试。

用法：
    .venv\Scripts\python.exe scripts/generate_sample_docs.py
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document as DocxDocument
from reportlab.pdfgen import canvas as pdf_canvas

EXAMPLES = Path(__file__).resolve().parent.parent / "examples"


def make_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("图神经网络综述", level=1)
    doc.add_paragraph("Graph Neural Network（GNN）是一种用于图结构数据的深度学习模型。")
    doc.add_paragraph("GCN 提出了基于谱域的图卷积方法，是 GNN 的代表工作之一。")
    doc.add_paragraph("GraphSAGE 通过邻居采样改进了 GCN 的可扩展性。")
    doc.add_paragraph("注意力机制被引入图神经网络，代表性工作为 GAT。")
    doc.save(str(path))


def make_pdf(path: Path) -> None:
    c = pdf_canvas.Canvas(str(path))
    c.setFont("Helvetica", 12)
    lines = [
        "Knowledge Graph Construction with LLM",
        "Large Language Models (LLM) can extract entities and relations from documents.",
        "The extracted triples form a knowledge graph for downstream analysis.",
        "Entity alignment merges duplicates across documents to keep the graph clean.",
    ]
    y = 760
    for line in lines:
        c.drawString(72, y, line)
        y -= 24
    c.save()


def main() -> None:
    EXAMPLES.mkdir(exist_ok=True)
    make_docx(EXAMPLES / "图神经网络综述.docx")
    make_pdf(EXAMPLES / "knowledge_graph_llm.pdf")
    print(f"样例已生成到 {EXAMPLES}")
    print("  - 图神经网络综述.docx")
    print("  - knowledge_graph_llm.pdf")


if __name__ == "__main__":
    sys.exit(main())
